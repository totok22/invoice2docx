#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Core invoice processing engine.
Extracted from invoice_word_builder/generate_invoice_docs.py for GUI integration.
"""

from __future__ import annotations

import csv
import json
import os
import re
import shutil
import xml.etree.ElementTree as ET
from copy import deepcopy
from dataclasses import dataclass, field
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP, InvalidOperation
from pathlib import Path
from typing import Callable

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.comments import Comment
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from pypdf import PdfReader


MONEY = Decimal("0.01")

STATUS_PASS = "通过"
STATUS_BLOCKED = "严重差异"
STATUS_WARNING = "需确认"

SUMMARY_HEADERS = ["序号", "发票号码", "购买方抬头", "购买方税号", "销售方", "发票总额", "数据来源", "校验状态", "问题说明"]
ITEM_HEADERS = ["序号", "发票号码", "发票物资名称", "实际物资名称", "品名", "单位", "数量", "单价", "金额", "销售方", "存放地点"]
ISSUE_HEADERS = ["级别", "序号", "发票号码", "问题说明", "期望金额", "当前明细合计", "差额", "来源"]

ProgressCallback = Callable[[str, float], None]

@dataclass
class Item:
    invoice_item_name: str
    actual_name: str
    unit: str
    quantity: Decimal | None
    total: Decimal
    seller: str
    invoice_no: str
    source_file: str
    spec: str = ""
    storage_location: str = "工训楼"
    product_name: str = ""

    @property
    def unit_price(self) -> Decimal:
        if self.quantity and self.quantity != 0:
            return self.total / self.quantity
        return self.total


@dataclass
class Invoice:
    key: int
    invoice_no: str
    seller: str
    total: Decimal
    source: str
    items: list[Item] = field(default_factory=list)
    status: str = STATUS_WARNING
    issue: str = ""
    buyer_name: str = ""
    buyer_tax_id: str = ""


@dataclass
class Issue:
    level: str
    message: str
    key: int | str = ""
    invoice_no: str = ""
    expected: Decimal | None = None
    actual: Decimal | None = None
    source: str = ""

    @property
    def difference(self) -> Decimal | None:
        if self.expected is None or self.actual is None:
            return None
        return self.expected - self.actual


def d(value: str | int | float | Decimal | None) -> Decimal:
    if value is None:
        return Decimal("0")
    if isinstance(value, Decimal):
        return value
    text = str(value).replace(",", "").replace("¥", "").replace("￥", "").strip()
    if not text:
        return Decimal("0")
    try:
        return Decimal(text)
    except InvalidOperation:
        return Decimal("0")


def money(value: Decimal) -> Decimal:
    return value.quantize(MONEY, rounding=ROUND_HALF_UP)


def fmt_money(value: Decimal, currency: bool = False) -> str:
    prefix = "¥" if currency else ""
    return f"{prefix}{money(value):.2f}"


def fmt_optional_money(value: Decimal | None) -> str:
    return "" if value is None else fmt_money(value)


def fmt_decimal(value: Decimal | None, places: int = 8) -> str:
    if value is None:
        return ""
    quant = Decimal("1." + "0" * places)
    text = f"{value.quantize(quant, rounding=ROUND_HALF_UP):f}"
    if "." in text:
        text = text.rstrip("0").rstrip(".")
    return text


def clean_item_name(name: str) -> str:
    return re.sub(r"\*[^*]+\*", "", name).strip()


def invoice_key_from_name(path: Path) -> int:
    match = re.match(r"(\d+)\+", path.name)
    if match:
        return int(match.group(1))
    return 10_000


def invoice_key_issue(path: Path) -> Issue | None:
    if re.match(r"(\d+)\+", path.name):
        return None
    return Issue(
        level="warning",
        key="",
        message=f"文件命名不规范，无法从文件名识别序号：{path.name}",
        source=str(path),
    )


def child_text(elem: ET.Element, tag: str) -> str:
    found = elem.find(".//" + tag)
    return found.text.strip() if found is not None and found.text else ""


def pick_text(data: dict, *keys: str) -> str:
    for key in keys:
        value = data.get(key)
        if value not in (None, ""):
            return str(value).strip()
    return ""


def parse_xml_invoice(path: Path, key: int) -> Invoice:
    root = ET.parse(path).getroot()
    invoice_no = child_text(root, "EIid")
    seller = child_text(root, "SellerName")
    buyer_name = child_text(root, "BuyerName")
    buyer_tax_id = child_text(root, "BuyerIdNum")
    total = d(child_text(root, "TotalTax-includedAmount"))
    invoice = Invoice(key=key, invoice_no=invoice_no, seller=seller, total=total, source=str(path), buyer_name=buyer_name, buyer_tax_id=buyer_tax_id)

    parents = [elem for elem in root.iter() if elem.find("ItemName") is not None]
    last: Item | None = None
    for elem in parents:
        raw_name = child_text(elem, "ItemName")
        spec = child_text(elem, "SpecMod")
        unit = child_text(elem, "MeaUnits")
        quantity_text = child_text(elem, "Quantity")
        line_total = d(child_text(elem, "Amount")) + d(child_text(elem, "ComTaxAm"))

        if last and not quantity_text and raw_name == last.invoice_item_name:
            last.total += line_total
            continue

        quantity = d(quantity_text) if quantity_text else None
        item = Item(
            invoice_item_name=raw_name,
            actual_name=clean_item_name(raw_name),
            unit=unit,
            quantity=quantity,
            total=line_total,
            seller=seller,
            invoice_no=invoice_no,
            source_file=path.name,
            spec=spec,
        )
        invoice.items.append(item)
        last = item

    return invoice


MANUAL_PDF_ITEMS = {
    26: [("*电工仪器仪表*直流稳压电源", "台", "1", "128.62")],
    27: [("*塑料制品*3D打印机线材 黑色63100", "卷", "1", "200.13")],
    36: [("*计算机配套产品*多普瑞串口线", "条", "1", "32.00")],
    37: [
        ("*计算机网络设备*有人口红4g dtu模块cat1透传232/485转4gdt", "台", "1", "103.00"),
        ("*计算机网络设备*价外费用", "台", "1", "9.00"),
    ],
    71: [("*配电控制设备*公牛（BULL）立式USB插座 过载保护插排/插线板/接线板/排插 8位总控+2A2C 全长1.6米 GNV-UL1008U黑色", "个", "1", "107.00")],
}


def invoice_needs_better_items(invoice: Invoice) -> bool:
    if not invoice.items:
        return True
    item_sum = sum((item.total for item in invoice.items), Decimal("0"))
    return money(invoice.total - item_sum) != Decimal("0.00")


def pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def first_pdf_invoice_item_name(path: Path, actual_name: str) -> str:
    text = pdf_text(path)
    for line in text.splitlines():
        line = line.strip()
        if not line.startswith("*"):
            continue
        if actual_name in line:
            return line[: line.index(actual_name) + len(actual_name)]
        match = re.match(r"(\*[^*]+\*[^ ]+)", line)
        if match:
            return match.group(1)
    return actual_name


def extract_pdf_buyer(lines: list[str]) -> tuple[str, str]:
    tax_id_pattern = re.compile(r"[0-9A-Z]{15,20}")
    for idx, line in enumerate(lines):
        if re.fullmatch(r"\d{4}年\d{2}月\d{2}日", line) and idx + 2 < len(lines):
            maybe_name = lines[idx + 1].strip()
            maybe_tax_id = lines[idx + 2].strip()
            if maybe_name and tax_id_pattern.fullmatch(maybe_tax_id):
                return maybe_name, maybe_tax_id
    for idx, line in enumerate(lines):
        if len(line) >= 30 and re.fullmatch(r"[0-9A-Z]+", line):
            buyer_tax_id = line[-18:]
            if tax_id_pattern.fullmatch(buyer_tax_id):
                name_line = lines[idx - 1].strip() if idx > 0 else ""
                buyer_name = ""
                for suffix in ["有限公司", "公司"]:
                    if suffix in name_line and not name_line.endswith(suffix):
                        buyer_name = name_line.split(suffix, 1)[1].strip()
                        break
                if buyer_name:
                    return buyer_name, buyer_tax_id
    for idx, line in enumerate(lines):
        if tax_id_pattern.fullmatch(line):
            prev = lines[idx - 1].strip() if idx > 0 else ""
            if prev and "有限公司" not in prev and "公司" not in prev and not re.fullmatch(r"[一-龥]{2,4}", prev):
                return prev, line
    return "", ""


def split_combined_pdf_party_line(lines: list[str], buyer_name: str, buyer_tax_id: str) -> tuple[str, str]:
    if not buyer_name or not buyer_tax_id:
        return "", ""

    for idx, line in enumerate(lines):
        if buyer_name in line and not line.startswith("*"):
            seller_name = line.replace(buyer_name, "", 1).strip()
            tax_line = lines[idx + 1].strip() if idx + 1 < len(lines) else ""
            if tax_line.endswith(buyer_tax_id):
                seller_tax_id = tax_line[: -len(buyer_tax_id)].strip()
                return seller_name, seller_tax_id

    return "", ""


def extract_pdf_seller(lines: list[str], buyer_tax_id: str, buyer_name: str) -> str:
    seller_from_combined, _seller_tax_id = split_combined_pdf_party_line(lines, buyer_name, buyer_tax_id)
    if seller_from_combined:
        return seller_from_combined

    for idx, line in enumerate(lines):
        if buyer_tax_id and line == buyer_tax_id and idx + 1 < len(lines):
            return lines[idx + 1]

    return ""


def parse_pdf_amount_tax_line(line: str) -> tuple[str, Decimal | None, Decimal, Decimal, int] | None:
    match = re.search(r"(?:\d+(?:\.\d+)?%)\s*([\u4e00-\u9fa5A-Za-z]+)\s+(-?\d+\.\d{2})\s+(-?\d+\.\d{2})", line)
    quantity: Decimal | None = None
    if not match:
        match = re.search(r"\s([\u4e00-\u9fa5A-Za-z]+)\s+(-?\d+(?:\.\d+)?)\s+-?\d+(?:\.\d+)?\s+(-?\d+\.\d{2})\s+(-?\d+\.\d{2})\d+(?:\.\d+)?%$", line)
        if not match:
            return None
        quantity = d(match.group(2))
        return match.group(1), quantity, d(match.group(3)), d(match.group(4)), match.start()
    return match.group(1), quantity, d(match.group(2)), d(match.group(3)), match.start()


def normalize_pdf_item_name(parts: list[str]) -> str:
    return re.sub(r"\s+", "", "".join(parts)).strip()


def parse_pdf_items_from_text(lines: list[str], invoice: Invoice, source_file: str) -> list[Item]:
    items: list[Item] = []
    current_name_parts: list[str] = []

    for line in lines:
        if line.startswith("*"):
            current_name_parts = [line]
            parsed = parse_pdf_amount_tax_line(line)
            if parsed:
                unit, quantity, amount_without_tax, tax_amount, name_end = parsed
                raw_name = normalize_pdf_item_name([line[:name_end]])
                items.append(
                    Item(
                        invoice_item_name=raw_name,
                        actual_name=clean_item_name(raw_name),
                        unit=unit,
                        quantity=quantity or Decimal("1"),
                        total=amount_without_tax + tax_amount,
                        seller=invoice.seller,
                        invoice_no=invoice.invoice_no,
                        source_file=source_file,
                    )
                )
                current_name_parts = []
            continue

        if not current_name_parts:
            continue

        parsed = parse_pdf_amount_tax_line(line)
        if parsed:
            unit, quantity, amount_without_tax, tax_amount, _name_end = parsed
            raw_name = normalize_pdf_item_name(current_name_parts)
            items.append(
                Item(
                    invoice_item_name=raw_name,
                    actual_name=clean_item_name(raw_name),
                    unit=unit,
                    quantity=quantity or Decimal("1"),
                    total=amount_without_tax + tax_amount,
                    seller=invoice.seller,
                    invoice_no=invoice.invoice_no,
                    source_file=source_file,
                )
            )
            current_name_parts = []
        elif not re.search(r"¥|合\s*计|价税合计|订单|购买时间|收款人|复核人|开票人", line):
            current_name_parts.append(line)

    return items


def parse_pdf_invoice(path: Path) -> Invoice:
    key = invoice_key_from_name(path)
    text = pdf_text(path)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    long_numbers = re.findall(r"\b\d{20}\b", text)
    invoice_no = long_numbers[0] if long_numbers else ""
    buyer_name, buyer_tax_id = extract_pdf_buyer(lines)

    seller = extract_pdf_seller(lines, buyer_tax_id, buyer_name)

    amounts = [d(x) for x in re.findall(r"¥\s*([0-9]+(?:\.[0-9]{2})?)", text)]
    total = max(amounts) if amounts else d(re.search(r"\+([0-9]+(?:\.[0-9]{2})?)\+发票", path.name).group(1))
    invoice = Invoice(key=key, invoice_no=invoice_no, seller=seller, total=total, source=str(path), buyer_name=buyer_name, buyer_tax_id=buyer_tax_id)
    invoice.items.extend(parse_pdf_items_from_text(lines, invoice, path.name))

    if not invoice.items and key in MANUAL_PDF_ITEMS:
        for raw_name, unit, quantity, total_text in MANUAL_PDF_ITEMS[key]:
            invoice.items.append(
                Item(
                    invoice_item_name=raw_name,
                    actual_name=clean_item_name(raw_name),
                    unit=unit,
                    quantity=d(quantity),
                    total=d(total_text),
                    seller=seller,
                    invoice_no=invoice_no,
                    source_file=path.name,
                )
            )
    return invoice


def parse_aliyun_ocr_invoice(path: Path, key: int) -> Invoice:
    if not os.environ.get("ALIBABA_CLOUD_ACCESS_KEY_ID") or not os.environ.get("ALIBABA_CLOUD_ACCESS_KEY_SECRET"):
        raise RuntimeError("未配置阿里云 OCR 密钥，请设置 ALIBABA_CLOUD_ACCESS_KEY_ID 和 ALIBABA_CLOUD_ACCESS_KEY_SECRET。")

    try:
        from alibabacloud_credentials.client import Client as CredentialClient
        from alibabacloud_darabonba_stream.client import Client as StreamClient
        from alibabacloud_ocr_api20210707 import models as ocr_api_20210707_models
        from alibabacloud_ocr_api20210707.client import Client as OcrClient
        from alibabacloud_tea_openapi import models as open_api_models
        from alibabacloud_tea_util import models as util_models
    except ImportError as exc:
        raise RuntimeError("未安装阿里云 OCR SDK，请安装 invoice_app/requirements.txt 中的阿里云 OCR 依赖。") from exc

    credential = CredentialClient()
    config = open_api_models.Config(credential=credential)
    config.endpoint = "ocr-api.cn-hangzhou.aliyuncs.com"
    client = OcrClient(config)
    request = ocr_api_20210707_models.RecognizeInvoiceRequest(body=StreamClient.read_from_file_path(str(path)))
    response = client.recognize_invoice_with_options(request, util_models.RuntimeOptions())
    if response.status_code != 200:
        raise RuntimeError(f"阿里云 OCR 返回状态码 {response.status_code}")

    payload = json.loads(response.body.data)
    data = payload.get("data", {})
    if not data:
        raise RuntimeError("阿里云 OCR 未返回发票 data 字段。")

    invoice_no = pick_text(data, "invoiceNumber", "invoiceNo", "invoiceCode")
    seller = pick_text(data, "sellerName", "seller")
    buyer_name = pick_text(data, "buyerName", "purchaserName", "buyer")
    buyer_tax_id = pick_text(data, "buyerTaxNo", "buyerTaxNumber", "buyerRegisterNum", "buyerIdNum", "purchaserTaxNumber")
    total = d(pick_text(data, "totalAmount", "amountWithTax", "totalTaxIncludedAmount"))
    invoice = Invoice(
        key=key,
        invoice_no=invoice_no,
        seller=seller,
        total=total,
        source=f"{path} (Aliyun OCR)",
        buyer_name=buyer_name,
        buyer_tax_id=buyer_tax_id,
    )

    last: Item | None = None
    for detail in data.get("invoiceDetails", []) or []:
        raw_name = pick_text(detail, "itemName", "name", "goodsName")
        if not raw_name:
            continue
        unit = pick_text(detail, "unit", "unitName")
        quantity_text = pick_text(detail, "quantity", "qty")
        line_total = d(pick_text(detail, "amount", "withoutTaxAmount")) + d(pick_text(detail, "tax", "taxAmount"))

        if last and not quantity_text and raw_name == last.invoice_item_name:
            last.total += line_total
            continue

        item = Item(
            invoice_item_name=raw_name,
            actual_name=clean_item_name(raw_name),
            unit=unit,
            quantity=d(quantity_text) if quantity_text else None,
            total=line_total,
            seller=seller,
            invoice_no=invoice_no,
            source_file=f"{path.name} / Aliyun OCR",
        )
        invoice.items.append(item)
        last = item

    return invoice


def collect_invoices(invoice_dir: Path, ocr_mode: str = "off", progress: ProgressCallback | None = None) -> tuple[list[Invoice], list[Issue]]:
    pdfs = sorted(invoice_dir.glob("**/*发票.pdf"), key=lambda p: invoice_key_from_name(p))
    xmls = sorted(invoice_dir.glob("**/*.xml"))
    issues: list[Issue] = []

    for pdf in pdfs:
        issue = invoice_key_issue(pdf)
        if issue:
            issues.append(issue)

    pdfs_by_key: dict[int, list[Path]] = {}
    for pdf in pdfs:
        pdfs_by_key.setdefault(invoice_key_from_name(pdf), []).append(pdf)
    for key, key_pdfs in pdfs_by_key.items():
        if key != 10_000 and len(key_pdfs) > 1:
            issues.append(Issue(level="warning", key=key, message="同一序号下有多个发票 PDF，请确认是否重复或需要拆分。", source="; ".join(str(p) for p in key_pdfs)))

    xml_by_no: dict[str, Path] = {}
    for path in xmls:
        root = ET.parse(path).getroot()
        xml_by_no[child_text(root, "EIid")] = path

    invoices: list[Invoice] = []
    used_xmls: set[Path] = set()
    total_pdfs = len(pdfs)
    for i, pdf in enumerate(pdfs):
        if progress:
            progress(f"解析发票 {pdf.name}", (i + 1) / max(total_pdfs, 1) * 0.6)
        key = invoice_key_from_name(pdf)
        try:
            text = pdf_text(pdf)
        except Exception as exc:
            issues.append(Issue(level="error", key=key, message=f"PDF 转文字失败：{exc}", source=str(pdf)))
            continue
        invoice_no_match = re.search(r"发票号码[:：]?\s*(\d{20})", text)
        invoice_no = invoice_no_match.group(1) if invoice_no_match else ""
        if invoice_no and invoice_no in xml_by_no:
            invoice = parse_xml_invoice(xml_by_no[invoice_no], key)
            invoice.source = f"{xml_by_no[invoice_no]} (matched {pdf.name})"
            for item in invoice.items:
                item.source_file = f"{xml_by_no[invoice_no].name} / {pdf.name}"
            if invoice.items:
                invoice.items[0].invoice_item_name = first_pdf_invoice_item_name(pdf, invoice.items[0].actual_name)
            invoices.append(invoice)
            used_xmls.add(xml_by_no[invoice_no])
        else:
            try:
                invoice = parse_pdf_invoice(pdf)
                if ocr_mode == "always" or (ocr_mode == "auto" and invoice_needs_better_items(invoice)):
                    try:
                        invoice = parse_aliyun_ocr_invoice(pdf, key)
                    except Exception as exc:
                        issues.append(
                            Issue(
                                level="warning" if invoice.items and ocr_mode == "auto" else "error",
                                key=key,
                                invoice_no=invoice.invoice_no,
                                message=f"阿里云 OCR 未能完成解析：{exc}",
                                source=str(pdf),
                            )
                        )
                invoices.append(invoice)
            except Exception as exc:
                issues.append(Issue(level="error", key=key, invoice_no=invoice_no, message=f"PDF 明细未能自动解析：{exc}", source=str(pdf)))

    for xml in xmls:
        if xml not in used_xmls:
            invoices.append(parse_xml_invoice(xml, invoice_key_from_name(xml)))

    return sorted(invoices, key=lambda inv: (inv.key, inv.invoice_no)), issues


def validate_invoices(invoices: list[Invoice], issues: list[Issue], expected_buyer_name: str = "", expected_buyer_tax_id: str = "") -> list[Issue]:
    validation_issues = list(issues)
    buyer_names = {inv.buyer_name.strip() for inv in invoices if inv.buyer_name.strip()}

    if len(buyer_names) > 1:
        validation_issues.append(Issue(level="error", message="同一发票文件夹中出现多个购买方抬头，需要拆分文件夹或修正发票。", source="; ".join(sorted(buyer_names))))

    expected_buyer_name = expected_buyer_name.strip() or next(iter(buyer_names), "")

    if expected_buyer_name:
        mismatched = sorted(name for name in buyer_names if name != expected_buyer_name)
        if mismatched:
            validation_issues.append(Issue(level="error", message=f"存在与期望购买方抬头不一致的发票。期望：{expected_buyer_name}", source="; ".join(mismatched)))

    for invoice in invoices:
        item_sum = sum((item.total for item in invoice.items), Decimal("0"))
        difference = money(invoice.total - item_sum)

        if difference == Decimal("0.00"):
            invoice.status = STATUS_PASS
            invoice.issue = ""
        else:
            invoice.status = STATUS_BLOCKED
            invoice.issue = f"发票总额与明细合计相差 {fmt_money(difference)}，需要修正。"
            validation_issues.append(Issue("error", invoice.issue, invoice.key, invoice.invoice_no, invoice.total, item_sum, invoice.source))

        if not invoice.items:
            invoice.status = STATUS_BLOCKED
            invoice.issue = "没有可用物品明细。"
            validation_issues.append(Issue("error", invoice.issue, invoice.key, invoice.invoice_no, invoice.total, Decimal("0"), invoice.source))

        buyer_issues = []
        if not invoice.buyer_name:
            buyer_issues.append("缺少购买方抬头")
        elif expected_buyer_name and invoice.buyer_name != expected_buyer_name:
            buyer_issues.append(f"购买方抬头为 {invoice.buyer_name}，与本批 {expected_buyer_name} 不一致")
        if buyer_issues:
            invoice.status = STATUS_BLOCKED
            invoice.issue = "；".join(filter(None, [invoice.issue, "；".join(buyer_issues)]))
            validation_issues.append(Issue("error", "；".join(buyer_issues), invoice.key, invoice.invoice_no, source=invoice.source))

    return validation_issues


def clone_row(table, template_row):
    new_tr = deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def clear_data_rows(table, keep_rows: int = 1):
    for row in list(table.rows)[keep_rows:]:
        table._tbl.remove(row._tr)


def set_cell(cell, text: str):
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = run.font.name


def ensure_template_shape(template: Path, min_table_columns: int, label: str):
    doc = Document(template)
    if not doc.tables:
        raise RuntimeError(f"{label}模板缺少表格：{template}")
    if len(doc.tables[0].rows) < 2:
        raise RuntimeError(f"{label}模板至少需要表头行和样例数据行：{template}")
    if len(doc.tables[0].rows[0].cells) < min_table_columns:
        raise RuntimeError(f"{label}模板表格列数不足：{template}")


def update_reimburse_doc(invoices: list[Invoice], total: Decimal, out_path: Path, template: Path, document_date: str):
    ensure_template_shape(template, 3, "报账说明")
    shutil.copyfile(template, out_path)
    doc = Document(out_path)
    doc.paragraphs[0].text = (
        f"机械与车辆学院申请支出{fmt_money(total)} 元。方程式车队比赛物资采买。"
        f"人民币{fmt_money(total)}元需打款至学生账户如下："
    )
    doc.paragraphs[-1].text = document_date
    table = doc.tables[0]
    template_row = deepcopy(table.rows[1])
    clear_data_rows(table, 1)

    for invoice in invoices:
        first = invoice.items[0]
        suffix = "等" if len(invoice.items) > 1 else ""
        row = clone_row(table, template_row)
        values = [first.actual_name + suffix, first.invoice_item_name + suffix, fmt_money(invoice.total, currency=True)]
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)

    doc.save(out_path)


def update_acceptance_doc(items: list[Item], out_path: Path, template: Path, document_date: str, storage_location: str):
    ensure_template_shape(template, 9, "验收单")
    shutil.copyfile(template, out_path)
    doc = Document(out_path)
    doc.paragraphs[1].text = f"单位   机械与车辆学院                                                                                       {document_date}"
    table = doc.tables[0]
    template_row = deepcopy(table.rows[1])
    attachment_text = table.rows[1].cells[-1].text
    clear_data_rows(table, 1)

    for item in items:
        row = clone_row(table, template_row)
        values = [
            item.product_name or item.actual_name,
            item.unit,
            fmt_decimal(item.quantity),
            fmt_decimal(item.unit_price),
            fmt_money(item.total),
            item.seller,
            item.invoice_no,
            item.storage_location or storage_location,
            attachment_text,
        ]
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)

    doc.save(out_path)


def style_sheet(ws, frozen_row: int = 1):
    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.freeze_panes = f"A{frozen_row + 1}"
    ws.auto_filter.ref = ws.dimensions
    for column_cells in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in column_cells)
        width = min(max(max_len + 2, 10), 42)
        ws.column_dimensions[get_column_letter(column_cells[0].column)].width = width
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def mark_operator_columns(ws, editable_headers: dict[str, str], hidden_headers: set[str]):
    header_to_column = {cell.value: cell.column for cell in ws[1] if cell.value}
    editable_fill = PatternFill("solid", fgColor="FFF2CC")
    hidden_fill = PatternFill("solid", fgColor="E7E6E6")

    for header, note in editable_headers.items():
        column = header_to_column.get(header)
        if not column:
            continue
        letter = get_column_letter(column)
        ws[f"{letter}1"].fill = editable_fill
        ws[f"{letter}1"].font = Font(color="000000", bold=True)
        ws[f"{letter}1"].comment = Comment(note, "Invoice Word Builder")
        for cell in ws[letter][1:]:
            cell.fill = editable_fill

    for header in hidden_headers:
        column = header_to_column.get(header)
        if not column:
            continue
        letter = get_column_letter(column)
        ws[f"{letter}1"].fill = hidden_fill
        ws.column_dimensions[letter].hidden = True


def guide_intermediate_workbook(summary, detail, issue_sheet):
    mark_operator_columns(
        summary,
        {
            "购买方抬头": "只有解析错误或本批发票抬头不一致时才需要修改。",
            "销售方": "只有销售方解析错误或缺失时才需要修改。",
            "发票总额": "只有发票总额解析错误时才需要修改。",
            "问题说明": "查看阻断原因；无需手工填写。",
        },
        {"购买方税号", "数据来源"},
    )
    mark_operator_columns(
        detail,
        {
            "实际物资名称": "报账说明使用此列；默认已从发票物资名称中去掉分类星号。",
            "品名": "验收单使用此列；通常只在需要更规范名称时修改。",
            "单位": "只有解析错误或缺失时才需要修改。",
            "数量": "只有解析错误或缺失时才需要修改。",
            "金额": "含税金额；金额校验失败时优先检查此列。",
            "销售方": "验收单使用此列；只有解析错误或缺失时才需要修改。",
            "存放地点": "默认使用设置中的存放地点；不同物品地点不同时再修改。",
        },
        {"发票号码", "发票物资名称", "单价"},
    )
    mark_operator_columns(issue_sheet, {"问题说明": "按这里的提示修正“发票汇总”或“物品明细”。"}, set())
    summary.sheet_properties.tabColor = "4472C4"
    detail.sheet_properties.tabColor = "70AD47"
    issue_sheet.sheet_properties.tabColor = "ED7D31"


def write_intermediate_xlsx(invoices: list[Invoice], output_dir: Path, issues: list[Issue], storage_location: str, file_name: str = "invoice_intermediate.xlsx") -> Path:
    workbook_path = output_dir / file_name
    wb = Workbook()
    summary = wb.active
    summary.title = "发票汇总"
    summary.append(SUMMARY_HEADERS)
    for inv in invoices:
        summary.append([inv.key, inv.invoice_no, inv.buyer_name, inv.buyer_tax_id, inv.seller, float(money(inv.total)), inv.source, inv.status, inv.issue])

    detail = wb.create_sheet("物品明细")
    detail.append(ITEM_HEADERS)
    for inv in invoices:
        for item in inv.items:
            detail.append([
                inv.key, item.invoice_no, item.invoice_item_name, item.actual_name,
                item.product_name or item.actual_name, item.unit,
                float(item.quantity) if item.quantity is not None else "",
                float(money(item.unit_price)), float(money(item.total)),
                item.seller, item.storage_location or storage_location,
            ])

    issue_sheet = wb.create_sheet("问题清单")
    issue_sheet.append(ISSUE_HEADERS)
    for issue in issues:
        issue_sheet.append([
            issue.level, issue.key, issue.invoice_no, issue.message,
            float(money(issue.expected)) if issue.expected is not None else "",
            float(money(issue.actual)) if issue.actual is not None else "",
            float(money(issue.difference)) if issue.difference is not None else "",
            issue.source,
        ])

    for ws in wb.worksheets:
        style_sheet(ws)
    for ws in [summary, detail, issue_sheet]:
        for row in ws.iter_rows(min_row=2):
            status = row[5].value if ws.title == "发票汇总" and len(row) > 5 else None
            level = row[0].value if ws.title == "问题清单" else None
            fill = None
            if status == STATUS_PASS:
                fill = PatternFill("solid", fgColor="E2F0D9")
            elif level == "warning":
                fill = PatternFill("solid", fgColor="FFF2CC")
            elif status == STATUS_BLOCKED or level == "error":
                fill = PatternFill("solid", fgColor="F4CCCC")
            if fill:
                for cell in row:
                    cell.fill = fill

    guide_intermediate_workbook(summary, detail, issue_sheet)
    wb.save(workbook_path)
    return workbook_path


def read_intermediate_xlsx(path: Path) -> list[Invoice]:
    wb = load_workbook(path, data_only=True)
    if "发票汇总" not in wb.sheetnames or "物品明细" not in wb.sheetnames:
        raise RuntimeError('中间 XLSX 必须包含"发票汇总"和"物品明细"两个工作表。')

    summary_ws = wb["发票汇总"]
    detail_ws = wb["物品明细"]
    summary_headers = [cell.value for cell in summary_ws[1]]
    detail_headers = [cell.value for cell in detail_ws[1]]

    missing_summary = [h for h in ["序号", "发票号码", "销售方", "发票总额"] if h not in summary_headers]
    missing_detail = [h for h in ITEM_HEADERS if h not in detail_headers]
    if missing_summary or missing_detail:
        raise RuntimeError(f"中间 XLSX 表头不完整。汇总缺少：{missing_summary}；明细缺少：{missing_detail}")

    summary_idx = {h: summary_headers.index(h) for h in summary_headers if h}
    detail_idx = {h: detail_headers.index(h) for h in detail_headers if h}

    invoices_by_key: dict[tuple[str, str], Invoice] = {}
    for row in summary_ws.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        key = row[summary_idx["序号"]]
        invoice_no = str(row[summary_idx["发票号码"]] or "")
        seller = str(row[summary_idx["销售方"]] or "")
        total = d(row[summary_idx["发票总额"]])
        source = str(row[summary_idx.get("数据来源", 0)] or path.name)
        buyer_name = str(row[summary_idx["购买方抬头"]] or "") if "购买方抬头" in summary_idx else ""
        buyer_tax_id = str(row[summary_idx["购买方税号"]] or "") if "购买方税号" in summary_idx else ""
        map_key = (str(key), invoice_no)
        invoices_by_key[map_key] = Invoice(key=int(key or 10_000), invoice_no=invoice_no, seller=seller, total=total, source=f"{path.name} / {source}", buyer_name=buyer_name, buyer_tax_id=buyer_tax_id)

    for row in detail_ws.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        key = row[detail_idx["序号"]]
        invoice_no = str(row[detail_idx["发票号码"]] or "")
        map_key = (str(key), invoice_no)
        invoice = invoices_by_key.get(map_key)
        if invoice is None:
            invoice = Invoice(key=int(key or 10_000), invoice_no=invoice_no, seller=str(row[detail_idx["销售方"]] or ""), total=Decimal("0"), source=str(path))
            invoices_by_key[map_key] = invoice
        raw_name = str(row[detail_idx["发票物资名称"]] or "")
        actual_name = str(row[detail_idx["实际物资名称"]] or clean_item_name(raw_name))
        item = Item(
            invoice_item_name=raw_name,
            actual_name=actual_name,
            unit=str(row[detail_idx["单位"]] or ""),
            quantity=d(row[detail_idx["数量"]]) if row[detail_idx["数量"]] not in (None, "") else None,
            total=d(row[detail_idx["金额"]]),
            seller=str(row[detail_idx["销售方"]] or invoice.seller),
            invoice_no=invoice_no,
            source_file=path.name,
            storage_location=str(row[detail_idx["存放地点"]] or "工训楼"),
            product_name=str(row[detail_idx["品名"]] or actual_name),
        )
        invoice.items.append(item)

    return sorted(invoices_by_key.values(), key=lambda inv: (inv.key, inv.invoice_no))


def write_audit(invoices: list[Invoice], items: list[Item], total: Decimal, output_dir: Path, issues: list[Issue], storage_location: str):
    invoice_csv = output_dir / "invoice_audit.csv"
    item_csv = output_dir / "acceptance_items.csv"
    json_path = output_dir / "invoice_audit.json"

    with invoice_csv.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["序号", "发票号码", "购买方抬头", "购买方税号", "销售方", "发票总额", "明细合计", "差额", "校验状态", "问题说明", "来源"])
        for inv in invoices:
            item_sum = sum((item.total for item in inv.items), Decimal("0"))
            writer.writerow([inv.key, inv.invoice_no, inv.buyer_name, inv.buyer_tax_id, inv.seller, fmt_money(inv.total), fmt_money(item_sum), fmt_money(inv.total - item_sum), inv.status, inv.issue, inv.source])

    with item_csv.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["发票号码", "品名", "单位", "数量", "单价", "金额", "来源或出售厂商", "存放地点", "来源文件"])
        for item in items:
            writer.writerow([
                item.invoice_no,
                item.product_name or item.actual_name,
                item.unit,
                fmt_decimal(item.quantity),
                fmt_decimal(item.unit_price),
                fmt_money(item.total),
                item.seller,
                item.storage_location or storage_location,
                item.source_file,
            ])

    payload = {
        "total": fmt_money(total),
        "invoice_count": len(invoices),
        "item_count": len(items),
        "has_blocking_issues": any(issue.level == "error" for issue in issues),
        "issues": [{"level": issue.level, "key": issue.key, "invoice_no": issue.invoice_no, "message": issue.message, "expected": fmt_optional_money(issue.expected), "actual": fmt_optional_money(issue.actual), "difference": fmt_optional_money(issue.difference), "source": issue.source} for issue in issues],
    }
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _safe_output_name(text: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]+", "_", text).strip()
    cleaned = re.sub(r"\s+", "_", cleaned)
    return cleaned or "run"


def make_run_output_dir(base_output_dir: Path, invoice_dir: Path, from_xlsx: Path | None) -> Path:
    source_name = from_xlsx.stem if from_xlsx else invoice_dir.name
    prefix = _safe_output_name(source_name)
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    candidate = base_output_dir / f"{prefix}-{timestamp}"
    counter = 2
    while candidate.exists():
        candidate = base_output_dir / f"{prefix}-{timestamp}-{counter}"
        counter += 1
    candidate.mkdir(parents=True, exist_ok=False)
    return candidate


@dataclass
class RunConfig:
    invoice_dir: Path
    reimburse_template: Path
    acceptance_template: Path
    output_dir: Path
    from_xlsx: Path | None = None
    storage_location: str = "工训楼"
    document_date: str = "2026年5月13日"
    expected_buyer_name: str = ""
    expected_buyer_tax_id: str = ""
    ocr_mode: str = "off"
    reimburse_name: str = "报账说明.docx"
    acceptance_name: str = "验收单.docx"
    allow_risky_generate: bool = False


@dataclass
class RunResult:
    success: bool
    invoices: list[Invoice]
    items: list[Item]
    total: Decimal
    issues: list[Issue]
    has_blocking_issues: bool
    intermediate_xlsx: Path | None = None
    reimburse_doc: Path | None = None
    acceptance_doc: Path | None = None
    error: str = ""


def run_pipeline(config: RunConfig, progress: ProgressCallback | None = None) -> RunResult:
    config.output_dir.mkdir(parents=True, exist_ok=True)
    run_output_dir = make_run_output_dir(config.output_dir, config.invoice_dir, config.from_xlsx)

    if progress:
        progress("读取发票数据…", 0.05)

    if config.from_xlsx:
        invoices = read_intermediate_xlsx(config.from_xlsx)
        initial_issues: list[Issue] = []
    else:
        invoices, initial_issues = collect_invoices(config.invoice_dir, config.ocr_mode, progress)

    if progress:
        progress("校验发票数据…", 0.65)

    items = [item for invoice in invoices for item in invoice.items]
    total = sum((invoice.total for invoice in invoices), Decimal("0"))
    issues = validate_invoices(invoices, initial_issues, config.expected_buyer_name, config.expected_buyer_tax_id)

    if progress:
        progress("导出审计文件…", 0.75)

    write_audit(invoices, items, total, run_output_dir, issues, config.storage_location)
    workbook_path = write_intermediate_xlsx(invoices, run_output_dir, issues, config.storage_location)

    has_blocking = any(issue.level == "error" for issue in issues)

    if has_blocking and not config.allow_risky_generate:
        return RunResult(
            success=False,
            invoices=invoices,
            items=items,
            total=total,
            issues=issues,
            has_blocking_issues=True,
            intermediate_xlsx=workbook_path,
            error="存在校验错误，Word 生成已跳过。请修正中间表后重新生成。",
        )

    if progress:
        progress("生成 Word 文档…", 0.85)

    reimburse_out = run_output_dir / config.reimburse_name
    acceptance_out = run_output_dir / config.acceptance_name
    update_reimburse_doc(invoices, total, reimburse_out, config.reimburse_template, config.document_date)
    update_acceptance_doc(items, acceptance_out, config.acceptance_template, config.document_date, config.storage_location)

    if progress:
        progress("完成", 1.0)

    return RunResult(
        success=True,
        invoices=invoices,
        items=items,
        total=total,
        issues=issues,
        has_blocking_issues=has_blocking,
        intermediate_xlsx=workbook_path,
        reimburse_doc=reimburse_out,
        acceptance_doc=acceptance_out,
    )
